from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "report_assets"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
GENERATED_DIR = ASSET_DIR / "generated"
OUTPUT_PATH = ROOT / "Portlify_Project_Report.docx"

PRIMARY = RGBColor(145, 87, 46)
TEXT = RGBColor(55, 40, 31)
MUTED = RGBColor(102, 89, 79)

TEAM_MEMBERS = [
    ("Soham Bhutkar", "253341038", "Led the project end-to-end, including planning, UI development, core implementation, integration, testing, report review, and presentation preparation"),
    ("Vishal Vishwakarma", "253341029", "Assisted in UI creation and motion graphics for marketing and presentation support"),
    ("Krushna Thombre", "253341013", "Assisted in backend-related work and selected features such as GitHub fetching and onboarding support"),
    ("Onkar Naik", "253341012", "Assisted in documentation, report structuring, and report writing with AI-assisted drafting support"),
]

INDEX_ITEMS = [
    "Chapter 1: Introduction",
    "Chapter 2: Problem Statement",
    "Chapter 3: Objectives",
    "Chapter 4: Literature Survey / Existing System",
    "Chapter 5: System Analysis and Requirements",
    "Chapter 6: System Architecture",
    "Chapter 7: Database Design / ER-Style Description",
    "Chapter 8: Module Description",
    "Chapter 9: Implementation Details",
    "Chapter 10: Testing and Results",
    "Chapter 11: Conclusion and Future Scope",
    "Chapter 12: Bibliography and References",
]


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
    candidates = [
        windir / "Fonts" / ("arialbd.ttf" if bold else "arial.ttf"),
        windir / "Fonts" / ("calibrib.ttf" if bold else "calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: tuple[int, int, int]) -> None:
    draw.line([start, end], fill=color, width=5)
    x1, y1 = end
    arrow_size = 10
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        direction = 1 if end[0] >= start[0] else -1
        draw.polygon(
            [
                (x1, y1),
                (x1 - direction * arrow_size, y1 - arrow_size // 2),
                (x1 - direction * arrow_size, y1 + arrow_size // 2),
            ],
            fill=color,
        )
    else:
        direction = 1 if end[1] >= start[1] else -1
        draw.polygon(
            [
                (x1, y1),
                (x1 - arrow_size // 2, y1 - direction * arrow_size),
                (x1 + arrow_size // 2, y1 - direction * arrow_size),
            ],
            fill=color,
        )


def draw_box(
    draw: ImageDraw.ImageDraw,
    coords: tuple[int, int, int, int],
    title: str,
    subtitle: str,
    fill: tuple[int, int, int],
    outline: tuple[int, int, int],
    title_font: ImageFont.ImageFont,
    body_font: ImageFont.ImageFont,
) -> None:
    draw.rounded_rectangle(coords, radius=24, fill=fill, outline=outline, width=4)
    x0, y0, _, _ = coords
    draw.text((x0 + 24, y0 + 18), title, font=title_font, fill=(40, 33, 27))
    body_y = y0 + 58
    for line in subtitle.split("\n"):
        draw.text((x0 + 24, body_y), line, font=body_font, fill=(70, 60, 52))
        body_y += 28


def build_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1700, 980), (248, 242, 235))
    draw = ImageDraw.Draw(img)
    title_font = load_font(30, bold=True)
    box_title = load_font(22, bold=True)
    body_font = load_font(18)
    draw.text((60, 32), "Portlify System Architecture", font=title_font, fill=(74, 48, 28))

    draw_box(
        draw,
        (110, 170, 460, 320),
        "Users and Browsers",
        "Students, job seekers,\nrecruiters, and reviewers",
        (255, 252, 249),
        (186, 149, 120),
        box_title,
        body_font,
    )
    draw_box(
        draw,
        (565, 170, 1115, 350),
        "Portlify Frontend",
        "React + TypeScript + Vite\nReact Router + React Query\nTailwind CSS + shadcn/ui + Framer Motion",
        (255, 252, 249),
        (186, 149, 120),
        box_title,
        body_font,
    )
    draw_box(
        draw,
        (1225, 170, 1585, 350),
        "External Services",
        "GitHub REST API\nGoogle Gemini API\nOpenAI API",
        (255, 252, 249),
        (186, 149, 120),
        box_title,
        body_font,
    )
    draw_box(
        draw,
        (505, 470, 1175, 700),
        "Supabase Platform",
        "Auth for login and session management\nPostgres tables for portfolio content\nEdge functions for imports, AI, and analytics",
        (255, 252, 249),
        (186, 149, 120),
        box_title,
        body_font,
    )
    draw_box(
        draw,
        (180, 760, 615, 900),
        "Publishing Layer",
        "Public route: /p/:username\nUnlisted route: /share/:token\nPrint route: /export/:mode",
        (255, 252, 249),
        (186, 149, 120),
        box_title,
        body_font,
    )
    draw_box(
        draw,
        (930, 760, 1470, 900),
        "Analytics and Feedback",
        "portfolio_views logs visitors\nlanding_feedback stores form responses\nDashboard summarizes progress and views",
        (255, 252, 249),
        (186, 149, 120),
        box_title,
        body_font,
    )

    arrow_color = (150, 101, 61)
    draw_arrow(draw, (460, 245), (565, 245), arrow_color)
    draw_arrow(draw, (1115, 245), (1225, 245), arrow_color)
    draw_arrow(draw, (840, 350), (840, 470), arrow_color)
    draw_arrow(draw, (760, 700), (430, 760), arrow_color)
    draw_arrow(draw, (940, 700), (1180, 760), arrow_color)
    draw_arrow(draw, (1225, 260), (1175, 550), arrow_color)
    draw_arrow(draw, (565, 265), (505, 550), arrow_color)

    img.save(path)


def build_er_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 1180), (248, 242, 235))
    draw = ImageDraw.Draw(img)
    title_font = load_font(30, bold=True)
    box_title = load_font(20, bold=True)
    body_font = load_font(16)
    draw.text((60, 32), "ER-Style Data Model of Portlify", font=title_font, fill=(74, 48, 28))

    central_fill = (255, 252, 249)
    outline = (186, 149, 120)

    draw_box(
        draw,
        (810, 115, 1110, 250),
        "profiles",
        "Stores user identity,\ncareer setup, and username",
        central_fill,
        outline,
        box_title,
        body_font,
    )
    draw_box(
        draw,
        (810, 340, 1110, 490),
        "portfolios",
        "Stores template, visibility,\nsection order, and sharing state",
        central_fill,
        outline,
        box_title,
        body_font,
    )

    child_boxes = {
        "bio_sections": (120, 610, 430, 730),
        "portfolio_projects": (470, 610, 850, 730),
        "skills": (900, 610, 1180, 730),
        "experiences": (1230, 610, 1560, 730),
        "education": (1600, 610, 1850, 730),
        "certifications": (120, 860, 470, 980),
        "contact_info": (530, 860, 840, 980),
        "custom_sections": (900, 860, 1240, 980),
        "portfolio_views": (1300, 860, 1600, 980),
        "landing_feedback": (1650, 860, 1850, 980),
    }
    subtitles = {
        "bio_sections": "One bio section per portfolio",
        "portfolio_projects": "Project title, links,\nproblem, solution, technologies",
        "skills": "Skill names, category,\nand proficiency metadata",
        "experiences": "Role, company, dates,\nand impact summary",
        "education": "Institution, degree,\nfield, year, and CGPA",
        "certifications": "Credential name, issuer,\nand dates",
        "contact_info": "Email and social links\nfor a portfolio",
        "custom_sections": "Extra title/body sections\nadded by the user",
        "portfolio_views": "Visitor analytics for\npublic and shared pages",
        "landing_feedback": "Landing page product\nfeedback submissions",
    }

    for name, coords in child_boxes.items():
        draw_box(draw, coords, name, subtitles[name], central_fill, outline, box_title, body_font)

    arrow_color = (150, 101, 61)
    draw_arrow(draw, (960, 250), (960, 340), arrow_color)
    draw_arrow(draw, (960, 490), (275, 610), arrow_color)
    draw_arrow(draw, (960, 490), (660, 610), arrow_color)
    draw_arrow(draw, (960, 490), (1040, 610), arrow_color)
    draw_arrow(draw, (960, 490), (1395, 610), arrow_color)
    draw_arrow(draw, (960, 490), (1725, 610), arrow_color)
    draw_arrow(draw, (960, 490), (285, 860), arrow_color)
    draw_arrow(draw, (960, 490), (685, 860), arrow_color)
    draw_arrow(draw, (960, 490), (1070, 860), arrow_color)
    draw_arrow(draw, (960, 490), (1450, 860), arrow_color)
    draw_arrow(draw, (1110, 180), (1650, 910), arrow_color)

    img.save(path)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading1.font.size = Pt(17)
    heading1.font.bold = True
    heading1.font.color.rgb = TEXT

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = TEXT

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = PRIMARY

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Times New Roman"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = MUTED


def add_centered_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False, color: RGBColor | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_footer_with_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        add_page_number(paragraph)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False, alignment: int | None = None) -> None:
    paragraph = doc.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_numbered_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(11)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cells[index].paragraphs[0]
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.0) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.name = "Times New Roman"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption_run.font.size = Pt(10.5)
    caption_run.font.color.rgb = MUTED


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    add_centered_paragraph(doc, "P.E.S.'s MODERN COLLEGE OF ARTS, SCIENCE & COMMERCE", size=16, bold=True)
    add_centered_paragraph(doc, "(AUTONOMOUS), GANESHKHIND, PUNE - 411016", size=14, bold=True)
    add_centered_paragraph(doc, "Academic Year 2025-26", size=13, bold=True)
    doc.add_paragraph()
    add_centered_paragraph(doc, "A PROJECT REPORT ON", size=14, bold=True, color=PRIMARY)
    add_centered_paragraph(doc, "PORTLIFY", size=24, bold=True, color=PRIMARY)
    doc.add_paragraph()
    add_centered_paragraph(doc, "Submitted in partial fulfillment of the requirements", size=12)
    add_centered_paragraph(doc, "for M.Sc. (Computer Science) Sem-II", size=12)
    doc.add_paragraph()
    add_centered_paragraph(doc, "Submitted By", size=13, bold=True)
    for name, roll_no, _ in TEAM_MEMBERS:
        add_centered_paragraph(doc, f"{name} - {roll_no}", size=12)
    doc.add_paragraph()
    add_centered_paragraph(doc, "Project Guide / Mentor", size=13, bold=True)
    add_centered_paragraph(doc, "Sonal Kulkarni Maam", size=12)
    doc.add_paragraph()
    add_centered_paragraph(doc, "Department of Computer Science", size=12)
    add_centered_paragraph(doc, "Modern College of Arts, Science & Commerce", size=12)


def add_certificate_page(doc: Document) -> None:
    doc.add_heading("Certificate", level=1)
    add_paragraph(
        doc,
        "This is to certify that Mr. Soham Bhutkar (Roll No. 253341038) and the accompanying project group members listed in this report have successfully completed the project titled 'Portlify' during the academic year 2025-26.",
    )
    add_paragraph(
        doc,
        "The work has been carried out under the guidance of Sonal Kulkarni Maam as part of the M.Sc. (Computer Science) Sem-II project work at P.E.S.'s Modern College of Arts, Science & Commerce (Autonomous), Ganeshkhind, Pune.",
    )
    add_paragraph(
        doc,
        "To the best of our knowledge, this report is an original academic submission based on the actual implementation and testing of the Portlify application.",
    )
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Project Guide / Mentor"
    table.rows[0].cells[1].text = "Head of Department"
    table.rows[1].cells[0].text = "\n\nSonal Kulkarni Maam"
    table.rows[1].cells[1].text = "\n\n________________________"


def add_acknowledgement_page(doc: Document) -> None:
    doc.add_heading("Acknowledgement", level=1)
    add_paragraph(
        doc,
        "We express our sincere gratitude to Sonal Kulkarni Maam for her guidance, encouragement, and valuable suggestions throughout the development of Portlify. Her support helped us convert our initial idea into a complete and presentable software project.",
    )
    add_paragraph(
        doc,
        "We are thankful to the Department of Computer Science and the faculty members of Modern College for providing the academic environment, technical direction, and continuous motivation required for this work.",
    )
    add_paragraph(
        doc,
        "We also acknowledge the efforts of all group members. Soham Bhutkar led the project implementation and overall coordination, while the remaining members supported the project through UI assistance, backend support, documentation, and presentation-related work. Their contributions helped complete the project in a collaborative manner.",
    )
    add_paragraph(
        doc,
        "Finally, we thank all users, open-source maintainers, and platform providers whose documentation and APIs made the implementation, testing, and refinement of Portlify possible.",
    )


def add_team_page(doc: Document) -> None:
    doc.add_heading("Group Details and Contribution Note", level=1)
    add_paragraph(
        doc,
        "This report is prepared as a group submission. The student names and currently available roll numbers are confirmed below. Contribution notes can still be refined later without changing the technical chapters of the report.",
    )
    add_table(
        doc,
        ["Sr. No.", "Member Name", "Roll Number", "Contribution"],
        [
            ["1", TEAM_MEMBERS[0][0], TEAM_MEMBERS[0][1], TEAM_MEMBERS[0][2]],
            ["2", TEAM_MEMBERS[1][0], TEAM_MEMBERS[1][1], TEAM_MEMBERS[1][2]],
            ["3", TEAM_MEMBERS[2][0], TEAM_MEMBERS[2][1], TEAM_MEMBERS[2][2]],
            ["4", TEAM_MEMBERS[3][0], TEAM_MEMBERS[3][1], TEAM_MEMBERS[3][2]],
        ],
    )
    add_paragraph(
        doc,
        "Before final academic submission, the contribution notes can be refined further without changing the remaining technical content of this report.",
        italic=True,
    )


def add_progress_page(doc: Document) -> None:
    doc.add_heading("Project Progress Report", level=1)
    add_table(
        doc,
        ["Month", "Duration", "Work Completed"],
        [
            ["January 2026", "2 weeks", "Topic selection, study of portfolio platforms, requirement gathering, and initial project planning"],
            ["February 2026", "4 weeks", "UI design, routing structure, authentication flow, onboarding, and database schema planning"],
            ["March 2026", "4 weeks", "Builder module, template system, GitHub import, LinkedIn parsing, AI features, analytics, and testing"],
            ["April 2026", "3 weeks", "Report writing, screenshot capture, result compilation, and presentation preparation"],
        ],
    )


def add_index_page(doc: Document) -> None:
    doc.add_heading("Index", level=1)
    add_table(
        doc,
        ["Sr. No.", "Name of Topic"],
        [[str(index), item] for index, item in enumerate(INDEX_ITEMS, start=1)],
    )


def add_chapter(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def build_report() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    build_architecture_diagram(GENERATED_DIR / "architecture.png")
    build_er_diagram(GENERATED_DIR / "er-diagram.png")

    doc = Document()
    configure_document(doc)

    add_cover_page(doc)
    page_break(doc)
    add_certificate_page(doc)
    page_break(doc)
    add_acknowledgement_page(doc)
    page_break(doc)
    add_team_page(doc)
    page_break(doc)
    add_progress_page(doc)
    page_break(doc)
    add_index_page(doc)
    page_break(doc)

    add_chapter(doc, "Chapter 1: Introduction")
    doc.add_heading("1.1 Project Overview", level=2)
    add_paragraph(
        doc,
        "Portlify is a full-stack portfolio-building platform created for students, freshers, job seekers, and professionals who want to present their work online without manually coding a personal website. The application combines guided onboarding, structured content editing, visual templates, portfolio sharing, and AI-assisted writing into one workflow.",
    )
    add_paragraph(
        doc,
        "The project is built with React, TypeScript, and Vite on the frontend, while Supabase provides authentication, database storage, and serverless edge functions. The result is a modern web application that supports both private editing and public portfolio publishing.",
    )
    doc.add_heading("1.2 Scope of the Project", level=2)
    add_bullets(
        doc,
        [
            "Allow users to create and manage one or more portfolios from a dashboard.",
            "Provide a guided onboarding process to capture career goals and suggested starter content.",
            "Enable editing of bio, projects, skills, experience, education, certifications, custom sections, contact information, and settings.",
            "Support public publishing, unlisted sharing, and print-friendly export.",
            "Offer AI-assisted content improvement and SOP generation from stored portfolio data.",
        ],
    )
    add_image(doc, SCREENSHOT_DIR / "landing-page.png", "Figure 1. Landing page of the Portlify application.", width=6.15)

    add_chapter(doc, "Chapter 2: Problem Statement")
    add_paragraph(
        doc,
        "Many students and early-career professionals know that a strong portfolio improves job and internship opportunities, but building such a portfolio usually requires web development skills, design effort, content planning, and hosting knowledge. As a result, many capable candidates either postpone making a portfolio or publish incomplete and low-quality profiles.",
    )
    add_paragraph(
        doc,
        "The problem addressed by Portlify is the absence of a simple, guided, and professional workflow that helps users build a portfolio quickly while still giving them control over content, presentation, and sharing. The platform is designed to reduce friction from signup to publishing and to provide tools that improve both content quality and visual presentation.",
    )

    add_chapter(doc, "Chapter 3: Objectives")
    add_bullets(
        doc,
        [
            "To build a secure account system using Supabase Auth.",
            "To provide guided onboarding for different user types and career goals.",
            "To support structured portfolio editing across major career sections.",
            "To provide five portfolio templates: Glass, Night Owl, Vibrant, Editorial, and Brutalist.",
            "To integrate GitHub public repository import for project seeding.",
            "To support LinkedIn PDF parsing for importing profile details.",
            "To add AI content polish for bio and project descriptions.",
            "To generate a Statement of Purpose from portfolio data.",
            "To publish portfolios through public and unlisted links.",
            "To capture view analytics and landing feedback for product insight.",
            "To provide print-ready export for portfolio or resume mode.",
        ],
    )

    add_chapter(doc, "Chapter 4: Literature Survey / Existing System")
    doc.add_heading("4.1 Existing Approaches", level=2)
    add_paragraph(
        doc,
        "Common approaches for creating a portfolio include coding a static website manually, using generic website builders, or sharing resume PDFs and social profiles. Manual coding provides flexibility but requires HTML, CSS, deployment, and maintenance knowledge. Generic site builders reduce coding effort but often lack project-specific structure for technical portfolios. Resume-only approaches are compact but do not offer live project presentation, analytics, or public portfolio routes.",
    )
    doc.add_heading("4.2 Observed Limitations", level=2)
    add_bullets(
        doc,
        [
            "Users spend too much time setting up structure before writing content.",
            "Generic tools do not always match recruiter-oriented portfolio needs.",
            "Project importing and portfolio analytics are often missing or fragmented.",
            "Content quality depends heavily on the user even when they need writing support.",
            "Private review and public publishing are usually handled separately.",
        ],
    )
    doc.add_heading("4.3 Why Portlify is Useful", level=2)
    add_paragraph(
        doc,
        "Portlify combines structured editing, template selection, public sharing, unlisted review links, analytics, AI writing support, GitHub import, LinkedIn import, and export in a single application. This directly addresses the limitations of fragmented or code-heavy portfolio workflows.",
    )

    add_chapter(doc, "Chapter 5: System Analysis and Requirements")
    doc.add_heading("5.1 Stakeholders", level=2)
    add_table(
        doc,
        ["Stakeholder", "Role in the System"],
        [
            ["Students / Job Seekers", "Create, update, publish, and export their professional portfolio"],
            ["Recruiters / Visitors", "View public portfolios and judge candidate work and presentation"],
            ["Private Reviewers", "Access unlisted links shared for feedback before publishing"],
            ["Project Team", "Develop, maintain, test, and enhance the Portlify platform"],
        ],
    )
    doc.add_heading("5.2 Functional Requirements", level=2)
    add_table(
        doc,
        ["Requirement ID", "Functional Requirement"],
        [
            ["FR1", "The system shall allow users to sign up, log in, and reset passwords."],
            ["FR2", "The system shall guide users through onboarding based on career stage and goal."],
            ["FR3", "The system shall create and manage multiple portfolios per user."],
            ["FR4", "The builder shall support editing of all major portfolio sections."],
            ["FR5", "The system shall allow template selection and preview before publishing."],
            ["FR6", "The system shall import project data from GitHub public repositories."],
            ["FR7", "The system shall parse LinkedIn PDF data into structured sections."],
            ["FR8", "The system shall publish portfolios through public and unlisted links."],
            ["FR9", "The system shall track portfolio views and display summary metrics."],
            ["FR10", "The system shall export a print-friendly version of the portfolio."],
            ["FR11", "The system shall offer AI-assisted content polish and SOP generation."],
        ],
    )
    doc.add_heading("5.3 Non-Functional Requirements", level=2)
    add_bullets(
        doc,
        [
            "Usability: The UI should remain simple enough for non-developers to follow.",
            "Performance: Core routes should load quickly and updates should feel responsive.",
            "Security: Authenticated content must remain protected using Supabase authentication and row-level access controls.",
            "Reliability: Portfolio data should be stored consistently and retrievable across sessions.",
            "Scalability: The structure should support multiple users, portfolios, and sections without redesigning the app.",
            "Maintainability: Modular components, hooks, routes, and edge functions should remain easy to extend.",
        ],
    )

    add_chapter(doc, "Chapter 6: System Architecture")
    add_paragraph(
        doc,
        "Portlify follows a modern client-platform architecture. The frontend is a single-page React application that handles navigation, user interactions, preview rendering, and data entry. Supabase acts as the backend platform by providing authentication, Postgres storage, and edge functions for external integrations and AI-based features.",
    )
    add_paragraph(
        doc,
        "Public visitors access published portfolio routes, while authenticated users work through protected routes such as onboarding, dashboard, builder, template selection, preview, export, and SOP generation. External APIs are called through controlled edge functions instead of directly exposing business logic in the client.",
    )
    add_bullets(
        doc,
        [
            "Browser users interact with the React frontend.",
            "Protected routes rely on Supabase session state and profile status.",
            "Custom hooks and React Query fetch and store structured portfolio data.",
            "Supabase Postgres stores profiles, portfolios, sections, analytics, and feedback.",
            "Edge functions connect the project to GitHub, Gemini, and OpenAI services.",
            "Published routes render the chosen template and log analytics to portfolio_views.",
        ],
    )
    add_image(doc, GENERATED_DIR / "architecture.png", "Figure 2. High-level system architecture of Portlify.", width=6.2)

    add_chapter(doc, "Chapter 7: Database Design / ER-Style Description")
    add_paragraph(
        doc,
        "The database design is centered around the portfolio entity. Each authenticated user has a profile and can manage one or more portfolios. Every portfolio stores template, visibility, and section configuration, while related tables store section-wise portfolio content.",
    )
    add_image(doc, GENERATED_DIR / "er-diagram.png", "Figure 3. ER-style data model used by Portlify.", width=6.2)
    add_table(
        doc,
        ["Table Name", "Purpose in the System"],
        [
            ["profiles", "Stores user identity, username, onboarding state, and career preferences"],
            ["portfolios", "Stores portfolio metadata such as template, visibility, and sharing configuration"],
            ["bio_sections", "Stores the main profile summary shown at the top of the portfolio"],
            ["portfolio_projects", "Stores project entries with description, links, and technologies"],
            ["skills", "Stores skill names, categories, and proficiency-related information"],
            ["experiences", "Stores work or internship experience details"],
            ["education", "Stores educational history, degree, and academic details"],
            ["certifications", "Stores credentials, issuers, and certification links"],
            ["contact_info", "Stores email and social/profile URLs for the portfolio owner"],
            ["custom_sections", "Stores extra sections created by the user beyond standard fields"],
            ["portfolio_views", "Stores portfolio view analytics for public and shared routes"],
            ["landing_feedback", "Stores user feedback captured from the landing page form"],
        ],
    )

    add_chapter(doc, "Chapter 8: Module Description")
    add_table(
        doc,
        ["Module", "Description"],
        [
            ["Authentication Module", "Handles signup, login, logout, password reset, and protected routing through Supabase Auth"],
            ["Onboarding Module", "Collects user type, goal, starter mode, and template preference to seed the first portfolio"],
            ["Dashboard Module", "Displays portfolio list, completion status, visibility state, and basic view metrics"],
            ["Builder Module", "Allows editing of bio, projects, skills, experience, education, certifications, contact, custom sections, and settings"],
            ["Template Module", "Lets the user select from five portfolio templates with different styles and layouts"],
            ["Public Sharing Module", "Publishes public username-based routes and unlisted share links"],
            ["Analytics Module", "Logs portfolio views and shows dashboard summary values"],
            ["GitHub Import Module", "Fetches public repositories and converts them into portfolio project entries"],
            ["LinkedIn Import Module", "Parses LinkedIn PDF content into structured profile data using AI"],
            ["AI Writing Module", "Polishes content and generates SOP text from actual portfolio data"],
            ["Export Module", "Builds a print-friendly document view and triggers browser print/export"],
        ],
    )
    doc.add_heading("8.1 Route Summary", level=2)
    add_table(
        doc,
        ["Route", "Access", "Purpose"],
        [
            ["/", "Public", "Landing page introducing the platform and collecting feedback"],
            ["/login, /signup", "Public", "Authentication entry points"],
            ["/forgot-password, /reset-password", "Public", "Password recovery flow"],
            ["/onboarding", "Protected", "Guided setup after account creation"],
            ["/dashboard", "Protected", "Portfolio overview and management"],
            ["/builder", "Protected", "Edit structured portfolio sections"],
            ["/preview", "Protected", "Preview current portfolio before publishing"],
            ["/templates", "Protected", "Switch between the five design templates"],
            ["/sop-generator", "Protected", "Generate a Statement of Purpose from stored profile data"],
            ["/export/:mode", "Protected", "Prepare print-friendly output for browser print or PDF export"],
            ["/p/:username", "Public", "Open the user's public portfolio"],
            ["/share/:token", "Public", "Open an unlisted portfolio using a share token"],
        ],
    )

    add_chapter(doc, "Chapter 9: Implementation Details")
    doc.add_heading("9.1 Technology Stack", level=2)
    add_table(
        doc,
        ["Layer", "Technology Used", "Purpose"],
        [
            ["Frontend", "React + TypeScript + Vite", "Single-page application and developer workflow"],
            ["Routing", "React Router", "Page routing for public and protected flows"],
            ["UI", "Tailwind CSS + shadcn/ui + Framer Motion", "Styling, components, and animations"],
            ["State / Data", "TanStack React Query + custom hooks", "Data fetching, caching, and updates"],
            ["Backend Platform", "Supabase", "Authentication, Postgres tables, and edge functions"],
            ["PDF Parsing Support", "pdfjs-dist", "Client-side text extraction support for LinkedIn PDF import"],
            ["Testing", "Vitest + Testing Library", "Unit and component tests"],
        ],
    )
    doc.add_heading("9.2 Edge Functions and Integrations", level=2)
    add_table(
        doc,
        ["Function / Integration", "Actual Role in Portlify"],
        [
            ["sync-github", "Calls the GitHub API and converts public repositories into project-ready data"],
            ["parse-linkedin", "Uses Google Gemini to extract structured experience, skills, education, and contact data from LinkedIn PDF text"],
            ["polish-content", "Uses Gemini to rewrite, shorten, strengthen, or suggest improvements for portfolio text"],
            ["generate-sop", "Uses OpenAI to generate a Statement of Purpose from portfolio content"],
            ["log-portfolio-view", "Logs hashed visitor information and share channel into portfolio_views"],
        ],
    )
    doc.add_heading("9.3 Important Implementation Notes", level=2)
    add_bullets(
        doc,
        [
            "Portfolios support private, public, and unlisted visibility states.",
            "Users can reorder or hide sections and can also mark some sections as not applicable.",
            "Template rendering is shared between preview and public portfolio routes.",
            "The export route prepares a print layout and then triggers the browser print dialog.",
            "The landing page includes a five-question feedback form whose data is stored in landing_feedback.",
        ],
    )
    doc.add_heading("9.4 Interface Screenshots", level=2)
    add_image(doc, SCREENSHOT_DIR / "onboarding-import-flow.png", "Figure 4. Guided onboarding with GitHub and LinkedIn import choices.", width=6.15)
    add_image(doc, SCREENSHOT_DIR / "dashboard.png", "Figure 5. Dashboard after guided setup with completion and visibility summary.", width=6.15)
    add_image(doc, SCREENSHOT_DIR / "builder.png", "Figure 6. Builder page for editing the bio section and AI-assisted writing controls.", width=6.15)
    add_image(doc, SCREENSHOT_DIR / "template-selection.png", "Figure 7. Template selection page showing all five available designs.", width=6.15)
    add_image(doc, SCREENSHOT_DIR / "public-portfolio.png", "Figure 8. Public portfolio page rendered from the selected template.", width=6.15)

    add_chapter(doc, "Chapter 10: Testing and Results")
    doc.add_heading("10.1 Automated Verification", level=2)
    add_table(
        doc,
        ["Verification", "Observed Result"],
        [
            ["Production build", "Passed using npm run build on 23 March 2026"],
            ["Unit and component tests", "Passed: 9 test files and 20 tests"],
            ["Bundle creation", "dist output generated successfully by Vite build"],
            ["Warnings observed", "Large chunk warning and pdf.js eval warning appeared, but build still completed successfully"],
        ],
    )
    doc.add_heading("10.2 Manual Functional Checks", level=2)
    add_table(
        doc,
        ["Test Case", "Result", "Observation"],
        [
            ["Landing page load", "Passed", "Application loads correctly and exposes features, templates, workflow, and feedback form"],
            ["Account creation", "Passed", "Signup created an account and moved the user into onboarding"],
            ["Onboarding flow", "Passed", "Step-based setup completed successfully and seeded a starter portfolio"],
            ["Dashboard access", "Passed", "Dashboard displayed completion, visibility, and portfolio card details"],
            ["Builder access", "Passed", "Structured editing of section content was available after onboarding"],
            ["Template selection", "Passed", "Five portfolio templates were visible and selectable"],
            ["Public publishing", "Passed", "A username-based public route opened the selected portfolio"],
            ["Export route", "Code verified", "Protected export route and printable component are implemented in the application"],
        ],
    )
    doc.add_heading("10.3 Key Results", level=2)
    add_bullets(
        doc,
        [
            "The project successfully connects content editing, template switching, and public sharing in one workflow.",
            "The data model supports both structured sections and flexible custom sections.",
            "Live integrations for GitHub import, LinkedIn parsing, AI polish, and SOP generation are implemented as separate features.",
            "The system is suitable for recruiter-facing portfolio publishing as well as academic demonstration.",
        ],
    )

    add_chapter(doc, "Chapter 11: Conclusion and Future Scope")
    add_paragraph(
        doc,
        "Portlify successfully meets the objective of creating a guided portfolio-building platform for students and professionals. The project demonstrates how a modern React frontend can be combined with Supabase services and AI-backed edge functions to create a complete real-world product instead of a static academic prototype.",
    )
    add_paragraph(
        doc,
        "The application supports a full user journey from account creation and onboarding to editing, preview, publishing, sharing, analytics, and export. The actual implementation available in the repository matches the corrected documentation in this report and replaces the previously incorrect references to unrelated domains and technology stacks.",
    )
    doc.add_heading("11.1 Future Scope", level=2)
    add_bullets(
        doc,
        [
            "Add custom domain mapping for live portfolios.",
            "Improve analytics with trend charts, geography, and source-based reporting.",
            "Add richer collaboration support for mentors or peers reviewing a portfolio draft.",
            "Expand AI assistance for content suggestions across all editable sections.",
            "Support more advanced export layouts and downloadable resume variations.",
            "Introduce deeper portfolio personalization such as theme-level configuration and media galleries.",
        ],
    )

    add_chapter(doc, "Chapter 12: Bibliography and References")
    add_numbered_list(
        doc,
        [
            "Project source repository for Portlify, including routes, hooks, components, tests, and Supabase functions.",
            "React Documentation - https://react.dev/",
            "Vite Documentation - https://vite.dev/",
            "Supabase Documentation - https://supabase.com/docs",
            "Tailwind CSS Documentation - https://tailwindcss.com/docs",
            "shadcn/ui Documentation - https://ui.shadcn.com/",
            "TanStack React Query Documentation - https://tanstack.com/query/latest",
            "Framer Motion Documentation - https://www.framer.com/motion/",
            "GitHub REST API Documentation - https://docs.github.com/en/rest",
            "Google Gemini API Documentation - https://ai.google.dev/",
            "OpenAI API Documentation - https://platform.openai.com/docs",
        ],
    )

    add_footer_with_page_numbers(doc)
    doc.save(OUTPUT_PATH)
    print(f"Created report at: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
