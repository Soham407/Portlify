from __future__ import annotations

from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_portlify_report import (
    GENERATED_DIR,
    ROOT,
    SCREENSHOT_DIR,
    TEAM_MEMBERS,
    build_architecture_diagram,
    build_er_diagram,
)


OUTPUT_PATH = ROOT / "Portlify_Project_Report_OJT_Official.docx"

COLLEGE_NAME = "Modern College of Arts, Science and Commerce (Autonomous), Ganeshkhind, Pune-16."
COURSE_NAME = "M.Sc. (Computer Science)-I NEP-I"
ACADEMIC_YEAR = "2025-2026"
MENTOR_NAME = "Sonal Kulkarni Madam"
TRAINING_CONTEXT = "Portlify Product Development Project"
DEPARTMENT_NAME = "Department of Computer Science"

INDEX_ROWS = [
    ("CERTIFICATE", True),
    ("CERTIFICATE FROM ORGANIZATION / TRAINING UNIT", True),
    ("ACKNOWLEDGEMENT", True),
    ("INDEX", True),
    ("1. Abstract", False),
    ("2. Introduction of Organization / Training Context", False),
    ("3. Objectives of Training", False),
    ("4. Technologies / Tools Used", False),
    ("5. Work Done During Training", False),
    ("6. Weekly Progress", False),
    ("7. Skills Learned", False),
    ("8. Challenges Faced and Solutions", False),
    ("9. Results / Output Screenshots", False),
    ("10. Testing and Validation", False),
    ("11. Advantages and Limitations", False),
    ("12. Future Enhancements", False),
    ("13. Conclusion", False),
    ("14. References / Bibliography", False),
    ("15. Appendix", False),
]


def set_font(run, size: float = 12, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_page_margins(section) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)


def configure_styles(doc: Document) -> None:
    for section in doc.sections:
        set_page_margins(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_restart_page_numbering(section, start: int = 1, fmt: str | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    if fmt:
        pg_num_type.set(qn("w:fmt"), fmt)


def add_page_field(paragraph, *, roman: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, 11)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* ROMAN" if roman else "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure_footer(section, *, roman: bool = False) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_field(paragraph, roman=roman)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: float | None = 0.5,
    bold: bool = False,
    italic: bool = False,
    size: float = 12,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Inches(first_line_indent) if first_line_indent is not None else None
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)


def add_centered_line(doc: Document, text: str, *, size: float = 12, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.first_line_indent = None
        run = paragraph.add_run(item)
        set_font(run, 12)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.first_line_indent = None
        run = paragraph.add_run(item)
        set_font(run, 12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(header)
        set_font(run, size=11, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_font(run, size=10.5)
    doc.add_paragraph()


def add_image(doc: Document, image_path: Path, caption: str, *, width: float = 5.9) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.line_spacing = 1.0
    caption_para.paragraph_format.first_line_indent = None
    caption_run = caption_para.add_run(caption)
    set_font(caption_run, size=10.5, italic=True)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_centered_line(doc, "Progressive Education Society’s", size=13, bold=True)
    add_centered_line(doc, COLLEGE_NAME, size=14, bold=True)
    doc.add_paragraph()
    add_centered_line(doc, "ON-JOB TRAINING (OJT) REPORT", size=16, bold=True)
    doc.add_paragraph()
    add_centered_line(doc, "ON", size=12, bold=True)
    add_centered_line(doc, "PORTLIFY", size=18, bold=True)
    add_centered_line(doc, f"Training Context: {TRAINING_CONTEXT}", size=12)
    doc.add_paragraph()
    add_centered_line(doc, "Submitted By", size=13, bold=True)
    for name, roll_no, _ in TEAM_MEMBERS:
        add_centered_line(doc, f"{name} : {roll_no}", size=12)
    doc.add_paragraph()
    add_centered_line(doc, f"Project Guide / Mentor: {MENTOR_NAME}", size=12)
    add_centered_line(doc, DEPARTMENT_NAME, size=12)
    add_centered_line(doc, f"{COURSE_NAME} Academic Year: {ACADEMIC_YEAR}", size=12, bold=True)


def add_college_certificate(doc: Document) -> None:
    add_heading(doc, "CERTIFICATE", level=1)
    add_paragraph(
        doc,
        "This is to certify that Mr. Soham Bhutkar (Roll No. 253341038), along with the project group members Vishal Vishwakarma (253341029), Krushna Thombre (253341013), and Onkar Naik (253341012), has successfully completed the On-Job Training work through the development of the project titled “Portlify”.",
    )
    add_paragraph(
        doc,
        f"The work was carried out under the guidance of {MENTOR_NAME} in the {DEPARTMENT_NAME}, {COLLEGE_NAME} as part of the {COURSE_NAME} curriculum for the academic year {ACADEMIC_YEAR}.",
    )
    add_paragraph(
        doc,
        "During this period, the group worked on requirement analysis, interface design, frontend development, backend integration, AI-assisted features, testing, documentation, and presentation preparation related to the Portlify platform.",
    )
    add_paragraph(
        doc,
        "To the best of our knowledge, the work carried out by the group is genuine and satisfactory.",
    )
    doc.add_paragraph()
    add_paragraph(doc, "Place: Pune", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    add_paragraph(doc, "Date: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
    doc.add_paragraph()
    add_table(
        doc,
        ["Project Guide / Mentor", "Head of Department"],
        [["\n\n" + MENTOR_NAME, "\n\n________________________"]],
    )
    add_table(
        doc,
        ["Internal Examiner", "External Examiner"],
        [["\n\n________________________", "\n\n________________________"]],
    )


def add_organization_certificate_placeholder(doc: Document) -> None:
    add_heading(doc, "CERTIFICATE FROM ORGANIZATION / TRAINING UNIT", level=1)
    add_paragraph(
        doc,
        "This page is intentionally reserved for an organization-side training certificate if your department requires a separate signed document from the training organization, internship unit, or sponsoring authority.",
    )
    add_paragraph(
        doc,
        "No separate company certificate details were available at the time this official-format report was generated, so a placeholder page has been kept instead of adding incorrect information.",
        italic=True,
    )
    add_paragraph(doc, "Organization / Unit Name: ____________________________________________", first_line_indent=None)
    add_paragraph(doc, "Industry Mentor / Authority: _________________________________________", first_line_indent=None)
    add_paragraph(doc, "Training Duration: _________________________________________________", first_line_indent=None)
    add_paragraph(doc, "Authorized Signature and Seal: ______________________________________", first_line_indent=None)


def add_acknowledgement(doc: Document) -> None:
    add_heading(doc, "ACKNOWLEDGEMENT", level=1)
    add_paragraph(
        doc,
        f"We express our sincere gratitude to {MENTOR_NAME} for her guidance, encouragement, and timely feedback throughout the development of Portlify. Her support helped us shape the project into a complete and presentable OJT submission.",
    )
    add_paragraph(
        doc,
        f"We are thankful to the {DEPARTMENT_NAME} and the faculty members of {COLLEGE_NAME} for providing the academic environment, technical direction, and motivation required to complete this work.",
    )
    add_paragraph(
        doc,
        "We also acknowledge the contribution of all team members. Soham Bhutkar led the project end-to-end, Vishal Vishwakarma supported UI creation and motion graphics, Krushna Thombre supported backend-related features such as GitHub fetching and onboarding, and Onkar Naik contributed to documentation and report structuring.",
    )
    add_paragraph(
        doc,
        "Finally, we thank the maintainers of React, Supabase, Tailwind CSS, GitHub, Google Gemini, OpenAI, and other tools and libraries that supported the implementation and verification of Portlify.",
    )


def add_index_page(doc: Document) -> None:
    add_heading(doc, "INDEX", level=1)
    rows = [[str(index), title, ""] for index, (title, _) in enumerate(INDEX_ROWS, start=1)]
    add_table(doc, ["Sr. No.", "Title", "Page No."], rows)
    add_paragraph(
        doc,
        "The index is kept in table form to match the preferred college submission style. Page numbers are populated automatically after generation.",
        italic=True,
        first_line_indent=None,
    )


def populate_index_table_pages(doc_path: Path) -> None:
    entry_lines = []
    for title, roman in INDEX_ROWS:
        escaped = title.replace("'", "''")
        entry_lines.append(
            f"  [pscustomobject]@{{Title='{escaped}'; Roman={'$true' if roman else '$false'}}}"
        )

    script = f"""
$path = '{str(doc_path)}'
$entries = @(
{chr(10).join(entry_lines)}
)

function ConvertTo-Roman([int]$number) {{
  $map = @(
    [pscustomobject]@{{Value=1000; Symbol='M'}},
    [pscustomobject]@{{Value=900; Symbol='CM'}},
    [pscustomobject]@{{Value=500; Symbol='D'}},
    [pscustomobject]@{{Value=400; Symbol='CD'}},
    [pscustomobject]@{{Value=100; Symbol='C'}},
    [pscustomobject]@{{Value=90; Symbol='XC'}},
    [pscustomobject]@{{Value=50; Symbol='L'}},
    [pscustomobject]@{{Value=40; Symbol='XL'}},
    [pscustomobject]@{{Value=10; Symbol='X'}},
    [pscustomobject]@{{Value=9; Symbol='IX'}},
    [pscustomobject]@{{Value=5; Symbol='V'}},
    [pscustomobject]@{{Value=4; Symbol='IV'}},
    [pscustomobject]@{{Value=1; Symbol='I'}}
  )
  $result = ''
  foreach ($item in $map) {{
    while ($number -ge $item.Value) {{
      $result += $item.Symbol
      $number -= $item.Value
    }}
  }}
  return $result
}}

$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open($path)
  $doc.Fields.Update() | Out-Null

  $indexTable = $null
  foreach ($table in $doc.Tables) {{
    if ($table.Columns.Count -lt 3) {{ continue }}
    $h1 = ($table.Cell(1,1).Range.Text -replace '[\\r\\a]', '').Trim()
    $h2 = ($table.Cell(1,2).Range.Text -replace '[\\r\\a]', '').Trim()
    $h3 = ($table.Cell(1,3).Range.Text -replace '[\\r\\a]', '').Trim()
    if ($h1 -eq 'Sr. No.' -and $h2 -eq 'Title' -and $h3 -eq 'Page No.') {{
      $indexTable = $table
      break
    }}
  }}

  if (-not $indexTable) {{
    throw 'Index table not found.'
  }}

  $headingPages = @{{}}
  foreach ($para in $doc.Paragraphs) {{
    $text = ($para.Range.Text -replace '[\\r\\a]', '').Trim()
    if ([string]::IsNullOrWhiteSpace($text)) {{ continue }}
    $styleName = ''
    try {{
      $styleName = $para.Range.Style.NameLocal
    }} catch {{
      $styleName = ''
    }}
    if ($styleName -notlike 'Heading*') {{ continue }}

    foreach ($entry in $entries) {{
      if ($text -eq $entry.Title -and -not $headingPages.ContainsKey($text)) {{
        $page = [int]$para.Range.Information(1)
        if ($entry.Roman) {{
          $headingPages[$text] = ConvertTo-Roman $page
        }} else {{
          $headingPages[$text] = [string]$page
        }}
      }}
    }}
  }}

  for ($row = 2; $row -le $indexTable.Rows.Count; $row++) {{
    $titleText = ($indexTable.Cell($row, 2).Range.Text -replace '[\\r\\a]', '').Trim()
    $pageText = if ($headingPages.ContainsKey($titleText)) {{ $headingPages[$titleText] }} else {{ '' }}
    $cellRange = $indexTable.Cell($row, 3).Range
    $cellRange.End = $cellRange.End - 1
    $cellRange.Text = $pageText
  }}

  $doc.Save()
}} finally {{
  if ($doc) {{ $doc.Close() }}
  if ($word) {{ $word.Quit() }}
}}
"""

    subprocess.run(
        ["powershell", "-NoProfile", "-Command", script],
        check=True,
    )


def add_main_heading(doc: Document, text: str) -> None:
    add_heading(doc, text, level=1)


def build_report() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    build_architecture_diagram(GENERATED_DIR / "architecture.png")
    build_er_diagram(GENERATED_DIR / "er-diagram.png")

    doc = Document()
    configure_styles(doc)

    cover_section = doc.sections[0]
    set_page_margins(cover_section)

    add_cover_page(doc)

    front_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_margins(front_section)
    set_restart_page_numbering(front_section, start=1, fmt="upperRoman")
    configure_footer(front_section, roman=True)

    add_college_certificate(doc)
    page_break(doc)
    add_organization_certificate_placeholder(doc)
    page_break(doc)
    add_acknowledgement(doc)
    page_break(doc)
    add_index_page(doc)

    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_margins(main_section)
    set_restart_page_numbering(main_section, start=1, fmt="decimal")
    configure_footer(main_section, roman=False)

    add_main_heading(doc, "1. Abstract")
    add_paragraph(
        doc,
        "Portlify is a full-stack portfolio-building web application developed as the practical OJT work for this report. The system helps students, freshers, job seekers, and professionals create structured online portfolios without manually coding a personal website. It combines guided onboarding, project import, AI-assisted content improvement, public sharing, analytics, and export in a single platform.",
    )
    add_paragraph(
        doc,
        "The application is built with React, TypeScript, and Vite on the frontend, while Supabase provides authentication, database storage, and serverless edge functions. The completed project demonstrates requirement analysis, UI implementation, backend integration, external API usage, testing, and technical documentation in one end-to-end workflow.",
    )

    add_main_heading(doc, "2. Introduction of Organization / Training Context")
    add_paragraph(
        doc,
        "This OJT work was centered on the development of Portlify as a product-oriented software project in the academic environment. Since the work focused on building a complete application rather than working within a separate external company system, the training context is represented through the project itself and the software engineering workflow followed during implementation.",
    )
    add_paragraph(
        doc,
        "Portlify addresses a practical need in the career and portfolio space. Many students and early-career professionals have strong projects and skills but do not have the time, design support, or technical knowledge required to build a polished portfolio site. Portlify reduces that barrier by guiding users from account creation to publishing.",
    )
    add_paragraph(
        doc,
        "The training work involved problem understanding, interface design, frontend development, Supabase integration, AI-backed parsing and writing support, dashboard and preview flows, documentation, and result validation. This made the project a suitable OJT-style implementation with both technical depth and practical relevance.",
    )

    add_main_heading(doc, "3. Objectives of Training")
    add_bullets(
        doc,
        [
            "To understand the complete workflow of building and delivering a modern web application.",
            "To create a guided portfolio-building platform for students, freshers, and job seekers.",
            "To implement secure authentication, structured content management, and public portfolio publishing.",
            "To integrate GitHub import, LinkedIn PDF parsing, AI-assisted writing, and SOP generation.",
            "To practice real project activities such as requirement analysis, UI creation, debugging, testing, and report preparation.",
        ],
    )

    add_main_heading(doc, "4. Technologies / Tools Used")
    add_table(
        doc,
        ["Layer / Tool", "Technology", "Purpose"],
        [
            ["Frontend", "React + TypeScript + Vite", "Single-page application development"],
            ["Routing", "React Router", "Public and protected route navigation"],
            ["UI and Styling", "Tailwind CSS + shadcn/ui + Framer Motion", "Design, components, and motion"],
            ["Backend Platform", "Supabase", "Authentication, database, and edge functions"],
            ["Database", "Supabase Postgres", "Structured storage for profiles, portfolios, and related entities"],
            ["Data Fetching", "TanStack React Query + custom hooks", "Fetching, caching, and updates"],
            ["PDF Text Extraction", "pdfjs-dist", "Reads LinkedIn PDF text in the client"],
            ["AI Services", "Google Gemini and OpenAI", "LinkedIn parsing, content polishing, and SOP generation"],
            ["Testing", "Vitest + Testing Library", "Unit and component verification"],
            ["Version Control", "Git and GitHub", "Source code management"],
        ],
    )

    add_main_heading(doc, "5. Work Done During Training")
    add_heading(doc, "5.1 Requirement Analysis and Project Planning", level=2)
    add_paragraph(
        doc,
        "The initial phase focused on understanding the problem faced by students and job seekers when trying to present their work online. Based on this, the workflow was planned around signup, guided onboarding, structured editing, preview, publishing, and export.",
    )
    add_heading(doc, "5.2 Frontend and User Experience Development", level=2)
    add_paragraph(
        doc,
        "The frontend was developed using React and TypeScript. Core pages such as the landing page, login, signup, onboarding, dashboard, builder, preview, templates, export, and SOP generator were created and connected through route-based navigation.",
    )
    add_heading(doc, "5.3 Backend Integration and Data Modeling", level=2)
    add_paragraph(
        doc,
        "Supabase was used for authentication, user profiles, portfolio storage, contact information, projects, skills, education, experience, certifications, portfolio views, and landing feedback. The data model supports both structured sections and flexible custom sections.",
    )
    add_heading(doc, "5.4 Feature Implementation", level=2)
    add_bullets(
        doc,
        [
            "Guided onboarding for user type, career field, goals, starter content, and template preference.",
            "Builder support for bio, projects, skills, education, experience, certifications, contact, and custom sections.",
            "Five templates: Glass, Night Owl, Vibrant, Editorial, and Brutalist.",
            "GitHub repository fetching to seed project entries.",
            "LinkedIn PDF parsing using client-side extraction and AI-backed structuring.",
            "AI content polishing and SOP generation from portfolio data.",
            "Public portfolio route, unlisted sharing, analytics logging, and print-friendly export.",
        ],
    )
    add_image(doc, GENERATED_DIR / "architecture.png", "Figure 1. High-level system architecture of Portlify.")

    add_main_heading(doc, "6. Weekly Progress")
    add_table(
        doc,
        ["Week / Phase", "Focus Area", "Activities Completed"],
        [
            ["Week 1", "Planning and study", "Studied portfolio platform requirements, analyzed workflow, and finalized the Portlify project direction."],
            ["Week 2", "UI and core flow", "Built foundational UI, route structure, authentication flow, and onboarding screens."],
            ["Week 3", "Builder and templates", "Implemented dashboard, builder modules, preview flow, template selection, and structured portfolio sections."],
            ["Week 4", "Integration and validation", "Added GitHub fetching, LinkedIn PDF parsing, AI support, screenshots, testing, documentation, and presentation preparation."],
        ],
    )

    add_main_heading(doc, "7. Skills Learned")
    add_bullets(
        doc,
        [
            "React and TypeScript based application development",
            "Supabase authentication and database integration",
            "UI design and structured component-based development",
            "Handling external APIs and AI-backed workflows",
            "Testing, debugging, and feature verification",
            "Documentation, report writing, and presentation preparation",
            "Team coordination and module-wise contribution management",
        ],
    )

    add_main_heading(doc, "8. Challenges Faced and Solutions")
    add_table(
        doc,
        ["Challenge", "Observed Issue", "Solution Followed"],
        [
            ["Content structure", "Users needed guidance, not just blank fields.", "Introduced onboarding, starter content, and structured section ordering."],
            ["LinkedIn import", "LinkedIn data was available only through PDF export.", "Used pdfjs-dist for text extraction and an edge function with Gemini for structured parsing."],
            ["Portfolio flexibility", "Different users needed different section priorities.", "Added user-type and goal-based starter section order logic."],
            ["Preview personalization", "Users needed layout control without breaking templates.", "Implemented section layout selection and reusable template rendering logic."],
            ["Verification", "The documentation had to match the real repository features.", "Validated routes, tables, edge functions, screenshots, build output, and tests before finalizing the report."],
        ],
    )

    add_main_heading(doc, "9. Results / Output Screenshots")
    add_paragraph(
        doc,
        "The final application provides a working flow from onboarding to portfolio publishing. Screenshots below show the actual implemented pages used in the project.",
    )
    add_image(doc, SCREENSHOT_DIR / "landing-page.png", "Figure 2. Landing page of the Portlify application.", width=5.8)
    add_image(doc, SCREENSHOT_DIR / "onboarding-import-flow.png", "Figure 3. Guided onboarding with GitHub and LinkedIn import choices.", width=5.8)
    add_image(doc, SCREENSHOT_DIR / "dashboard.png", "Figure 4. Dashboard after guided setup.", width=5.8)
    add_image(doc, SCREENSHOT_DIR / "builder.png", "Figure 5. Builder page for editing portfolio content.", width=5.8)
    add_image(doc, SCREENSHOT_DIR / "template-selection.png", "Figure 6. Template selection page.", width=5.8)
    add_image(doc, SCREENSHOT_DIR / "public-portfolio.png", "Figure 7. Public portfolio output.", width=5.8)

    add_main_heading(doc, "10. Testing and Validation")
    add_heading(doc, "10.1 Automated Verification", level=2)
    add_table(
        doc,
        ["Verification", "Result"],
        [
            ["Production build", "Passed using npm run build"],
            ["Unit and component tests", "Passed: 9 files and 20 tests"],
            ["Route and feature validation", "Confirmed against the actual repository and screenshots"],
        ],
    )
    add_heading(doc, "10.2 Manual Functional Checks", level=2)
    add_table(
        doc,
        ["Area", "Status", "Observation"],
        [
            ["Authentication", "Passed", "Login, signup, and protected routing are implemented."],
            ["Onboarding", "Passed", "User type, goals, import intent, and template selection work in a guided flow."],
            ["Builder and preview", "Passed", "Structured editing and preview flows are present."],
            ["Sharing", "Passed", "Public and unlisted portfolio routes are implemented."],
            ["Export", "Verified", "Protected export route and printable output are available."],
        ],
    )

    add_main_heading(doc, "11. Advantages and Limitations")
    add_heading(doc, "11.1 Advantages", level=2)
    add_bullets(
        doc,
        [
            "Reduces the effort required to create a professional online portfolio.",
            "Combines structured content editing, templates, sharing, and export in one platform.",
            "Supports real integrations such as GitHub import and LinkedIn PDF parsing.",
            "Provides AI-assisted writing support and SOP generation.",
        ],
    )
    add_heading(doc, "11.2 Limitations", level=2)
    add_bullets(
        doc,
        [
            "Some advanced customization features are still limited to the current template system.",
            "External AI-backed features depend on service availability and configured API access.",
            "Analytics are currently summary-oriented and can be expanded further in future versions.",
        ],
    )

    add_main_heading(doc, "12. Future Enhancements")
    add_bullets(
        doc,
        [
            "Custom domain support for live portfolios",
            "More advanced analytics and trend reporting",
            "Richer media galleries and personalization options",
            "Expanded AI assistance across more editable sections",
            "Additional export styles and downloadable resume formats",
        ],
    )

    add_main_heading(doc, "13. Conclusion")
    add_paragraph(
        doc,
        "The OJT work carried out through Portlify demonstrates the end-to-end development of a practical web application that solves a real presentation problem for students and job seekers. The project integrates frontend design, backend services, AI support, and public portfolio delivery into one system.",
    )
    add_paragraph(
        doc,
        "This report presents the work in the college’s official structure while keeping the technical content aligned with the actual implementation. The final result is both academically presentable and technically credible.",
    )

    add_main_heading(doc, "14. References / Bibliography")
    add_numbered(
        doc,
        [
            "Project source repository for Portlify.",
            "React Documentation - https://react.dev/",
            "Vite Documentation - https://vite.dev/",
            "Supabase Documentation - https://supabase.com/docs",
            "Tailwind CSS Documentation - https://tailwindcss.com/docs",
            "shadcn/ui Documentation - https://ui.shadcn.com/",
            "TanStack React Query Documentation - https://tanstack.com/query/latest",
            "Framer Motion Documentation - https://www.framer.com/motion/",
            "GitHub REST API Documentation - https://docs.github.com/en/rest",
            "Google Gemini API Documentation - https://ai.google.dev/",
            "OpenAI API Documentation - https://platform.openai.com/docs",
        ],
    )

    add_main_heading(doc, "15. Appendix")
    add_heading(doc, "15.1 ER-Style Data Model", level=2)
    add_image(doc, GENERATED_DIR / "er-diagram.png", "Figure 8. ER-style data model used by Portlify.", width=5.8)
    add_heading(doc, "15.2 Team Contribution Note", level=2)
    add_table(
        doc,
        ["Member", "Roll Number", "Contribution"],
        [[name, roll, contribution] for name, roll, contribution in TEAM_MEMBERS],
    )

    doc.save(OUTPUT_PATH)
    populate_index_table_pages(OUTPUT_PATH)
    print(f"Created official-format report at: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
